# Copyright (C) 2021 - 2022 Advanced Micro Devices, Inc. All rights reserved.
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in
# all copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.  IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN
# THE SOFTWARE.
import subprocess
import pathlib

import docx

import perflib.docx_emf_patch


def pdf2emf(path: pathlib.Path):
    """Convert PDF to EMF."""
    pdf, svg, emf = str(path), str(path.with_suffix(".svg")), str(
        path.with_suffix(".emf"))
    subprocess.check_call(["pdf2svg", pdf, svg])
    # Older versions of inkscape use -M.
    #subprocess.check_call(["inkscape", svg, "-M", emf])
    subprocess.check_call(["inkscape", svg, "--export-filename", emf])
    return emf


def make_docx(figs, docdir, outdirs, secondtype=None):

    docdir = pathlib.Path(docdir)

    document = docx.Document()

    document.add_heading('rocFFT benchmarks', 0)

    # document.add_paragraph("Each data point represents the median of " + str(nsample) + " values, with error bars showing the 95% confidence interval for the median.  Transforms are " + precision + "-precision, forward, and in-place.")

    # if secondtype == "gflops":
    #     document.add_paragraph(gflopstext)
    # if secondtype == "efficiency":
    #     document.add_paragraph(efficiencytext)

    # specfilename = os.path.join(outdir, "specs.txt")
    # if os.path.isfile(specfilename):
    #     with open(specfilename, "r") as f:
    #         specs = f.read()
    #     for line in specs.split("\n"):
    #         document.add_paragraph(line)

    for fig in figs:
        emfname = pdf2emf(fig.filename)
        document.add_picture(emfname, width=docx.shared.Inches(6))
        document.add_paragraph(fig.caption)

    document.save(str(docdir / 'figs.docx'))
